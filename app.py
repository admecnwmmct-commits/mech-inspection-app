from flask import Flask, render_template, request, session, send_file, jsonify, redirect
import os, io, json, shutil
from datetime import datetime

# ── Dynamic checklist loader ──────────────────────────────────────────────────
BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
CHECKLIST_JSON = os.path.join(BASE_DIR, 'checklist.json')
DEFAULT_JSON   = os.path.join(BASE_DIR, 'checklist_default.json')

def load_sections():
    try:
        with open(CHECKLIST_JSON) as f:
            return json.load(f)
    except Exception:
        from checklist_data import SECTIONS as S
        return S

def save_sections(data):
    with open(CHECKLIST_JSON, 'w') as f:
        json.dump(data, f, indent=2)

SECTIONS = load_sections()

app = Flask(__name__)
app.secret_key   = os.environ.get("SECRET_KEY",  "mech-bcr-inspection-2026")
ADMIN_PASSWORD   = os.environ.get("ADMIN_PASSWORD", "bctrailway2026")

# ── helpers ───────────────────────────────────────────────────────────────────
def answered_items(responses):
    """Return only items that have a yes/no/obs answer."""
    return {k: v for k, v in responses.items() if v.get('answer') in ('yes','no','obs')}

def collect_copy_to(header, responses):
    """Auto-collect unique Action By names from all NO items."""
    seen = set()
    names = []
    for v in responses.values():
        if v.get('answer') == 'no' and v.get('action_by','').strip():
            for n in v['action_by'].split(','):
                n = n.strip()
                if n and n not in seen:
                    seen.add(n)
                    names.append(n)
    return names

# ── PDF ───────────────────────────────────────────────────────────────────────
def generate_pdf(header, responses, other_remarks='', extra_copy_to='', bw=False):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, HRFlowable, KeepTogether
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    if bw:
        navy   = colors.black
        saffron= colors.black
        steel  = colors.HexColor('#444444')
        lgray  = colors.HexColor('#f0f0f0')
        mgray  = colors.HexColor('#cccccc')
        green  = colors.black
        red    = colors.black
        white  = colors.white
        black  = colors.black
    else:
        navy   = colors.HexColor('#0a1628')
        saffron= colors.HexColor('#f4a014')
        steel  = colors.HexColor('#3a7bd5')
        lgray  = colors.HexColor('#f5f5f5')
        mgray  = colors.HexColor('#d0d0d0')
        green  = colors.HexColor('#16a34a')
        red    = colors.HexColor('#dc2626')
        white  = colors.white
        black  = colors.black
    W = 17*cm

    title_style   = ParagraphStyle('T',  fontSize=13, fontName='Helvetica-Bold', textColor=white,  alignment=TA_CENTER, spaceAfter=3)
    sub_style     = ParagraphStyle('S',  fontSize=8.5,fontName='Helvetica',      textColor=saffron,alignment=TA_CENTER)
    intro_style   = ParagraphStyle('I',  fontSize=10, fontName='Helvetica-Bold', textColor=navy,   alignment=TA_CENTER, spaceBefore=8, spaceAfter=8)
    head2_style   = ParagraphStyle('H2', fontSize=10, fontName='Helvetica-Bold', textColor=navy,   spaceBefore=10, spaceAfter=4)
    head3_style   = ParagraphStyle('H3', fontSize=8.5,fontName='Helvetica-Bold', textColor=steel,  spaceBefore=5, spaceAfter=3)
    body_style    = ParagraphStyle('B',  fontSize=8,  fontName='Helvetica',      textColor=black,  leading=11)
    remark_style  = ParagraphStyle('R',  fontSize=7.5,fontName='Helvetica-Oblique',textColor=colors.HexColor('#444444'), leading=10)
    action_style  = ParagraphStyle('A',  fontSize=7.5,fontName='Helvetica-Bold', textColor=colors.HexColor('#0a1628'), leading=10)
    small_style   = ParagraphStyle('SM', fontSize=8,  fontName='Helvetica',      textColor=colors.HexColor('#555555'))
    obs_style     = ParagraphStyle('OB', fontSize=9,  fontName='Helvetica',      textColor=black,  leading=14, spaceAfter=4)
    copy_style    = ParagraphStyle('CP', fontSize=8.5,fontName='Helvetica',      textColor=black,  leading=13)
    copy_b_style  = ParagraphStyle('CB', fontSize=8.5,fontName='Helvetica-Bold', textColor=navy,   leading=13)

    story = []

    # 1. Inspection line
    loc   = header.get('location','')
    date  = header.get('date','')
    name  = header.get('name','')
    desig = header.get('designation','')
    itype = header.get('insp_type','')
    intro_line = f"Inspection of <b>{loc}</b> on <b>{date}</b> by <b>{desig}</b>"
    if itype:
        intro_line += f"<br/>Type: {itype}"
    story.append(Paragraph(intro_line, intro_style))
    story.append(HRFlowable(width=W, thickness=1, color=saffron, spaceAfter=8))

    # 3. Score summary
    ans_resp  = answered_items(responses)
    total_yes = sum(1 for v in ans_resp.values() if v.get('answer') == 'yes')
    total_no  = sum(1 for v in ans_resp.values() if v.get('answer') == 'no')
    total_ans = total_yes + total_no
    pct = round(100 * total_yes / total_ans) if total_ans else 0

    score_data = [["Items Inspected", "Satisfactory", "Deficiencies", "Compliance %"],
                  [str(total_ans), str(total_yes), str(total_no), f"{pct}%"]]
    score_tbl = Table(score_data, colWidths=[W/4]*4)
    score_tbl.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0), steel),
        ('TEXTCOLOR', (0,0),(-1,0), white),
        ('FONTNAME',  (0,0),(-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',  (0,0),(-1,-1), 9),
        ('ALIGN',     (0,0),(-1,-1), 'CENTER'),
        ('VALIGN',    (0,0),(-1,-1), 'MIDDLE'),
        ('BACKGROUND',(0,1),(-1,1), lgray),
        ('GRID',      (0,0),(-1,-1), 0.5, mgray),
        ('TOPPADDING',(0,0),(-1,-1), 5),
        ('BOTTOMPADDING',(0,0),(-1,-1), 5),
        ('FONTNAME',  (0,1),(-1,1), 'Helvetica-Bold'),
        ('TEXTCOLOR', (1,1),(1,1), green),
        ('TEXTCOLOR', (2,1),(2,1), red),
        ('FONTSIZE',  (0,1),(-1,1), 11),
    ]))
    story.append(score_tbl)
    story.append(Spacer(1, 0.5*cm))

    # 4. Per-section tables — ONLY answered items
    for section in SECTIONS:
        sec_rows = []
        for sub in section['subsections']:
            sub_rows = []
            for item in sub['items']:
                resp = responses.get(item['id'], {})
                if resp.get('answer') not in ('yes','no','obs'):
                    continue
                ans     = resp['answer']
                remark  = resp.get('remark','')
                action  = resp.get('action_by','') if ans == 'no' else ''
                if ans == 'obs':
                    # Observation-only item — spans status col, no yes/no
                    sub_rows.append([
                        Paragraph(item['label'], body_style),
                        Paragraph('—', ParagraphStyle('an', fontSize=8, fontName='Helvetica', textColor=colors.HexColor('#888888'), alignment=TA_CENTER)),
                        Paragraph(remark or '', ParagraphStyle('obsrm', fontSize=8, fontName='Helvetica-Oblique', textColor=colors.HexColor('#5a4000'), leading=11)),
                        Paragraph('', action_style),
                    ])
                else:
                    ans_up    = ans.upper()
                    ans_color = green if ans == 'yes' else red
                    sub_rows.append([
                        Paragraph(item['label'], body_style),
                        Paragraph(f"<b>{ans_up}</b>", ParagraphStyle('an', fontSize=8, fontName='Helvetica-Bold', textColor=ans_color, alignment=TA_CENTER)),
                        Paragraph(remark or '', remark_style),
                        Paragraph(action or '', action_style),
                    ])
            if sub_rows:
                sec_rows.append((sub['title'], sub_rows))

        if not sec_rows:
            continue

        story.append(Paragraph(f"{section['icon']}  {section['title'].upper()}", head2_style))
        story.append(HRFlowable(width=W, thickness=1.5, color=saffron, spaceAfter=5))

        for sub_title, rows in sec_rows:
            story.append(Paragraph(sub_title, head3_style))
            tbl_data = [[
                Paragraph("Inspection Point", ParagraphStyle('hb', fontSize=8, fontName='Helvetica-Bold', textColor=white)),
                Paragraph("Status",           ParagraphStyle('hb2',fontSize=8, fontName='Helvetica-Bold', textColor=white, alignment=TA_CENTER)),
                Paragraph("Remarks",          ParagraphStyle('hb3',fontSize=8, fontName='Helvetica-Bold', textColor=white)),
                Paragraph("Action By",        ParagraphStyle('hb4',fontSize=8, fontName='Helvetica-Bold', textColor=white)),
            ]] + rows
            col_w = [8.4*cm, 1.6*cm, 3.8*cm, 3.2*cm]
            tbl = Table(tbl_data, colWidths=col_w, repeatRows=1)
            tbl.setStyle(TableStyle([
                ('BACKGROUND',    (0,0),(-1,0), navy),
                ('TEXTCOLOR',     (0,0),(-1,0), white),
                ('VALIGN',        (0,0),(-1,-1), 'TOP'),
                ('ROWBACKGROUNDS',(0,1),(-1,-1), [white, lgray]),
                ('GRID',          (0,0),(-1,-1), 0.4, mgray),
                ('TOPPADDING',    (0,0),(-1,-1), 4),
                ('BOTTOMPADDING', (0,0),(-1,-1), 4),
                ('LEFTPADDING',   (0,0),(-1,-1), 5),
                ('ALIGN',         (1,0),(1,-1), 'CENTER'),
                ('LINEBELOW',     (0,-1),(-1,-1), 1, steel),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 0.25*cm))

    # 5. Other Remarks / Observations
    if other_remarks and other_remarks.strip():
        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph("OTHER REMARKS / OBSERVATIONS", head2_style))
        story.append(HRFlowable(width=W, thickness=1, color=saffron, spaceAfter=6))
        story.append(Paragraph(other_remarks.strip(), obs_style))

    # 6. Copy To section
    auto_names = collect_copy_to(header, responses)
    extra_names = [n.strip() for n in extra_copy_to.split(',') if n.strip()] if extra_copy_to else []
    all_copy = auto_names + [n for n in extra_names if n not in auto_names]

    if all_copy:
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph("COPY TO", head2_style))
        story.append(HRFlowable(width=W, thickness=1, color=saffron, spaceAfter=6))
        for name in all_copy:
            story.append(Paragraph(f"• <b>{name}</b> — For information and necessary action.", copy_style))

    # 7. Signature block — shows designation
    story.append(Spacer(1, 0.8*cm))
    sig_data = [
        ["Date & Time of Inspection", desig],
        [datetime.now().strftime("%d-%m-%Y  %H:%M"), "\n\n" + "_"*30],
    ]
    sig_tbl = Table(sig_data, colWidths=[W/2, W/2])
    sig_tbl.setStyle(TableStyle([
        ('FONTNAME',     (0,0),(-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',     (0,0),(-1,-1), 8.5),
        ('ALIGN',        (0,0),(-1,-1), 'CENTER'),
        ('GRID',         (0,0),(-1,-1), 0.5, mgray),
        ('BACKGROUND',   (0,0),(-1,0), lgray),
        ('TOPPADDING',   (0,0),(-1,-1), 5),
        ('BOTTOMPADDING',(0,0),(-1,-1), 12),
    ]))
    story.append(sig_tbl)

    doc.build(story)
    buf.seek(0)
    return buf


# ── DOCX ──────────────────────────────────────────────────────────────────────
def generate_docx(header, responses, other_remarks='', extra_copy_to='', bw=False):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    hdr_bg  = 'ffffff' if bw else '0a1628'
    hdr_txt = RGBColor(0,0,0) if bw else RGBColor(0xff,0xff,0xff)
    yes_clr = RGBColor(0,0,0) if bw else RGBColor(0x16,0xa3,0x4a)
    no_clr  = RGBColor(0,0,0) if bw else RGBColor(0xdc,0x26,0x26)
    ttl_clr = RGBColor(0,0,0) if bw else RGBColor(0x0a,0x16,0x28)
    sub_clr = RGBColor(0,0,0) if bw else RGBColor(0x3a,0x7b,0xd5)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Cm(2)
        sec.top_margin = sec.bottom_margin = Cm(2)

    # Inspection line
    loc   = header.get('location','')
    date  = header.get('date','')
    name  = header.get('name','')
    desig = header.get('designation','')
    itype = header.get('insp_type','')
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_i = intro.add_run(f"Inspection of {loc} on {date} by {desig}")
    run_i.bold = True
    run_i.font.size = Pt(11)
    run_i.font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
    if itype:
        intro.add_run(f"\nType: {itype}").font.size = Pt(9)

    doc.add_paragraph()

    # Score
    ans_resp  = answered_items(responses)
    total_yes = sum(1 for v in ans_resp.values() if v.get('answer') == 'yes')
    total_no  = sum(1 for v in ans_resp.values() if v.get('answer') == 'no')
    total_ans = total_yes + total_no
    pct = round(100 * total_yes / total_ans) if total_ans else 0
    sc = doc.add_paragraph()
    sc.add_run('COMPLIANCE SUMMARY: ').bold = True
    sc.add_run(f'Items Inspected: {total_ans}  |  Satisfactory: {total_yes}  |  Deficiencies: {total_no}  |  Compliance: {pct}%')
    doc.add_paragraph()

    # Sections — only answered items
    for section in SECTIONS:
        has_content = False
        for sub in section['subsections']:
            for item in sub['items']:
                resp = responses.get(item['id'], {})
                if resp.get('answer') in ('yes','no'):
                    has_content = True
                    break
        if not has_content:
            continue

        doc.add_heading(f"{section['icon']} {section['title']}", level=1)
        for sub in section['subsections']:
            sub_rows = []
            for item in sub['items']:
                resp = responses.get(item['id'], {})
                if resp.get('answer') not in ('yes','no','obs'):
                    continue
                sub_rows.append((item, resp))
            if not sub_rows:
                continue

            doc.add_heading(sub['title'], level=2)
            tbl2 = doc.add_table(rows=1, cols=4)
            tbl2.style = 'Table Grid'
            hdr_cells = tbl2.rows[0].cells
            for ci, hd in enumerate(['Inspection Point', 'Status', 'Remarks', 'Action By']):
                hdr_cells[ci].text = hd
                hdr_cells[ci].paragraphs[0].runs[0].bold = True
                hdr_cells[ci].paragraphs[0].runs[0].font.size = Pt(8)
                set_cell_bg(hdr_cells[ci], hdr_bg)
                hdr_cells[ci].paragraphs[0].runs[0].font.color.rgb = hdr_txt

            for i, (item, resp) in enumerate(sub_rows, 1):
                ans_raw = resp['answer']
                remark  = resp.get('remark','')
                action  = resp.get('action_by','') if ans_raw == 'no' else ''
                row = tbl2.add_row().cells
                row[0].text = item['label']
                row[0].paragraphs[0].runs[0].font.size = Pt(8)
                if ans_raw == 'obs':
                    row[1].text = 'Obs.'
                    row[1].paragraphs[0].runs[0].font.size = Pt(8)
                    row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x88,0x88,0x88)
                    row[2].text = remark
                    row[2].paragraphs[0].runs[0].font.size = Pt(8)
                    row[2].paragraphs[0].runs[0].font.italic = True
                    row[3].text = ''
                else:
                    ans = ans_raw.upper()
                    row[1].text = ans
                    row[2].text = remark
                    row[3].text = action
                    for ci in range(1,4):
                        row[ci].paragraphs[0].runs[0].font.size = Pt(8)
                    run2 = row[1].paragraphs[0].runs[0]
                    run2.bold = True
                    if ans == 'YES':
                        run2.font.color.rgb = yes_clr
                    elif ans == 'NO':
                        run2.font.color.rgb = no_clr
                    if action:
                        row[3].paragraphs[0].runs[0].bold = True
                        row[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
                if i % 2 == 0:
                    for ci in range(4):
                        set_cell_bg(row[ci], 'f5f5f5')
            for row in tbl2.rows:
                for ci, w in enumerate([9.5, 1.8, 3.5, 3.0]):
                    row.cells[ci].width = Cm(w)
            doc.add_paragraph()

    # Other Remarks
    if other_remarks and other_remarks.strip():
        doc.add_heading('OTHER REMARKS / OBSERVATIONS', level=1)
        doc.add_paragraph(other_remarks.strip())

    # Copy To
    auto_names  = collect_copy_to(header, responses)
    extra_names = [n.strip() for n in extra_copy_to.split(',') if n.strip()] if extra_copy_to else []
    all_copy    = auto_names + [n for n in extra_names if n not in auto_names]
    if all_copy:
        doc.add_heading('COPY TO', level=1)
        for cname in all_copy:
            p = doc.add_paragraph()
            p.add_run(f"{cname}").bold = True
            p.add_run(" — For information and necessary action.")

    # Signature — designation
    doc.add_paragraph()
    sig_tbl = doc.add_table(rows=2, cols=2)
    sig_tbl.style = 'Table Grid'
    for ci, hd in enumerate(["Date & Time of Inspection", desig]):
        sig_tbl.rows[0].cells[ci].text = hd
        sig_tbl.rows[0].cells[ci].paragraphs[0].runs[0].bold = True
        sig_tbl.rows[0].cells[ci].paragraphs[0].runs[0].font.size = Pt(8)
    sig_tbl.rows[1].cells[0].text = datetime.now().strftime("%d-%m-%Y  %H:%M")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── WhatsApp ──────────────────────────────────────────────────────────────────
def generate_whatsapp(header, responses, other_remarks='', extra_copy_to='', bw=False):
    ans_resp  = answered_items(responses)
    total_yes = sum(1 for v in ans_resp.values() if v.get('answer') == 'yes')
    total_no  = sum(1 for v in ans_resp.values() if v.get('answer') == 'no')
    total_ans = total_yes + total_no
    pct = round(100 * total_yes / total_ans) if total_ans else 0

    loc   = header.get('location','')
    date  = header.get('date','')
    name  = header.get('name','')
    desig = header.get('designation','')

    lines = [
        f"*Inspection of {loc} on {date} by {name} ({desig})*",
        "",
        "*COMPLIANCE SUMMARY*",
        f"Items Inspected: {total_ans}  |  Satisfactory: {total_yes}  |  Deficiencies: {total_no}  |  Compliance: {pct}%",
        "",
        "*DEFICIENCIES NOTED:*",
    ]

    nc_count = 0
    for section in SECTIONS:
        for sub in section['subsections']:
            for item in sub['items']:
                resp = responses.get(item['id'], {})
                ans = resp.get('answer','')
                if ans == 'no':
                    nc_count += 1
                    remark = resp.get('remark','')
                    action = resp.get('action_by','')
                    line = f"• {item['label'][:80]}"
                    if remark: line += f"\n  Remark: {remark}"
                    if action: line += f"\n  Action By: {action}"
                    lines.append(line)
                elif ans == 'obs':
                    remark = resp.get('remark','')
                    if remark:
                        lines.append(f"📝 {item['label'][:70]}\n  Obs: {remark}")
    if nc_count == 0:
        lines.append("Nil — Full Compliance")

    if other_remarks and other_remarks.strip():
        lines += ["", "*OTHER REMARKS / OBSERVATIONS:*", other_remarks.strip()]

    auto_names  = collect_copy_to(header, responses)
    extra_names = [n.strip() for n in extra_copy_to.split(',') if n.strip()] if extra_copy_to else []
    all_copy    = auto_names + [n for n in extra_names if n not in auto_names]
    if all_copy:
        lines += ["", "*COPY TO:*"]
        for cname in all_copy:
            lines.append(f"• {cname} — For information and necessary action.")

    lines += ["", f"_{datetime.now().strftime('%d-%m-%Y %H:%M')}_"]
    return "\n".join(lines)


# ── Admin routes ──────────────────────────────────────────────────────────────
@app.route('/admin/login', methods=['GET','POST'])
def admin_login():
    if request.method == 'POST':
        if request.form.get('password') == ADMIN_PASSWORD:
            session['admin'] = True
            return redirect('/admin')
        return render_template('admin_login.html', error='Incorrect password')
    return render_template('admin_login.html', error=None)

@app.route('/admin')
def admin():
    if not session.get('admin'):
        return redirect('/admin/login')
    return render_template('admin.html')

@app.route('/admin/data')
def admin_data():
    if not session.get('admin'):
        return jsonify({'error':'unauthorized'}), 403
    return jsonify(load_sections())

@app.route('/admin/save', methods=['POST'])
def admin_save():
    if not session.get('admin'):
        return jsonify({'status':'error','message':'unauthorized'}), 403
    try:
        data = request.get_json()
        save_sections(data)
        global SECTIONS
        SECTIONS = data
        return jsonify({'status':'ok'})
    except Exception as e:
        return jsonify({'status':'error','message':str(e)})

@app.route('/admin/reset', methods=['POST'])
def admin_reset():
    if not session.get('admin'):
        return jsonify({'status':'error'}), 403
    try:
        from checklist_data import SECTIONS as S
        save_sections(S)
        global SECTIONS
        SECTIONS = S
        return jsonify({'status':'ok'})
    except Exception as e:
        return jsonify({'status':'error','message':str(e)})

@app.route('/admin/logout')
def admin_logout():
    session.pop('admin', None)
    return redirect('/')

# ── Main routes ────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/checklist')
def checklist():
    sections_json = json.dumps(SECTIONS)
    return render_template('checklist.html', sections_json=sections_json)

@app.route('/submit', methods=['POST'])
def submit():
    data = request.get_json()
    session['header']    = data.get('header', {})
    session['responses'] = data.get('responses', {})
    return jsonify({'status': 'ok'})

@app.route('/summary')
def summary():
    header    = session.get('header', {})
    responses = session.get('responses', {})
    ans_resp  = answered_items(responses)
    total_yes = sum(1 for v in ans_resp.values() if v.get('answer') == 'yes')
    total_no  = sum(1 for v in ans_resp.values() if v.get('answer') == 'no')
    total_ans = total_yes + total_no
    pct = round(100 * total_yes / total_ans) if total_ans else 0
    nc_items = []
    for section in SECTIONS:
        for sub in section['subsections']:
            for item in sub['items']:
                resp = responses.get(item['id'], {})
                if resp.get('answer') == 'no':
                    nc_items.append({
                        'label':   item['label'],
                        'remark':  resp.get('remark',''),
                        'section': section['title'],
                        'action_by': resp.get('action_by',''),
                        'type': 'no'
                    })
                elif resp.get('answer') == 'obs' and resp.get('remark','').strip():
                    nc_items.append({
                        'label':   item['label'],
                        'remark':  resp.get('remark',''),
                        'section': section['title'],
                        'action_by': '',
                        'type': 'obs'
                    })
    auto_copy = collect_copy_to(header, responses)
    return render_template('summary.html',
                           header=header,
                           total_yes=total_yes, total_no=total_no,
                           total_ans=total_ans, pct=pct,
                           nc_items=nc_items,
                           auto_copy=auto_copy)

@app.route('/download/pdf', methods=['POST'])
def download_pdf():
    header         = session.get('header', {})
    responses      = session.get('responses', {})
    other_remarks  = request.form.get('other_remarks', '')
    extra_copy_to  = request.form.get('extra_copy_to', '')
    bw   = request.form.get('bw','0') == '1'
    buf  = generate_pdf(header, responses, other_remarks, extra_copy_to, bw)
    wa   = generate_whatsapp(header, responses, other_remarks, extra_copy_to)
    session['whatsapp'] = wa
    fname = f"Inspection_{header.get('location','BCT').replace(' ','_')}_{header.get('date','').replace('-','')}.pdf"
    return send_file(buf, as_attachment=True, download_name=fname, mimetype='application/pdf')

@app.route('/download/docx', methods=['POST'])
def download_docx():
    header         = session.get('header', {})
    responses      = session.get('responses', {})
    other_remarks  = request.form.get('other_remarks', '')
    extra_copy_to  = request.form.get('extra_copy_to', '')
    bw   = request.form.get('bw','0') == '1'
    buf  = generate_docx(header, responses, other_remarks, extra_copy_to, bw)
    wa   = generate_whatsapp(header, responses, other_remarks, extra_copy_to)
    session['whatsapp'] = wa
    fname = f"Inspection_{header.get('location','BCT').replace(' ','_')}_{header.get('date','').replace('-','')}.docx"
    return send_file(buf, as_attachment=True, download_name=fname, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/whatsapp')
def whatsapp_page():
    wa = session.get('whatsapp', '')
    return render_template('whatsapp.html', whatsapp=wa)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
